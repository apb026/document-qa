import streamlit as st
from google import genai
from io import BytesIO
from docx import Document
import re
from PIL import Image, ImageDraw, ImageFont
import io

# Function to add styled text in the Word document
def add_styled_text(doc, text, style=None, is_bold=False, is_italic=False, is_code=False):
    para = doc.add_paragraph()
    if is_code:
        run = para.add_run(text)
        run.font.name = 'Courier New'  # Use monospace font for code
        para.style = 'Normal'
    else:
        if style:
            para.style = style
        run = para.add_run(text)
        
        # Apply bold or italic formatting if needed
        if is_bold:
            run.bold = True
        if is_italic:
            run.italic = True

# Function to parse markdown-like syntax to Word formatting
def parse_markdown_to_word(doc, text):
    # Convert headers (e.g., # Header) to Heading 1 style
    text = re.sub(r'^\#\s+(.*)', lambda m: add_styled_text(doc, m.group(1), style='Heading 1'), text, flags=re.MULTILINE)
    
    # Convert bold text (e.g., **bold**) to actual bold
    while '**' in text:
        start = text.find('**')
        end = text.find('**', start + 2)
        if start == -1 or end == -1:
            break
        bold_text = text[start + 2:end]
        add_styled_text(doc, text[:start])
        add_styled_text(doc, bold_text, is_bold=True)
        text = text[end + 2:]
    
    # Add remaining text
    if text:
        add_styled_text(doc, text)

# Function to add a tip or warning box
def add_tip_or_warning_box(doc, text, box_type='tip'):
    para = doc.add_paragraph()
    para.add_run(text).italic = True
    
    if box_type == 'tip':
        para.paragraph_format.alignment = 1  # Center-aligned
        doc.add_paragraph("TIP:", style='Heading 3')
        para = doc.add_paragraph()
        para.add_run(text).bold = True
    elif box_type == 'warning':
        para.paragraph_format.alignment = 1
        doc.add_paragraph("WARNING:", style='Heading 3')
        para = doc.add_paragraph()
        para.add_run(text).bold = True

# Function to set persona
def set_persona(persona_description, query):
    return f"{persona_description}\n\n{query}"

# Function to generate the few-shot prompt
def get_few_shot_prompt(user_query, examples):
    example_text = "\n\n".join(examples)
    return f"Here are some examples of code documentation:\n\n{example_text}\n\nNow, for the following code:\n{user_query}"

# Function for RAG-based prompt generation
def rag_prompt(query, document_text):
    return f"Based on the document text below, answer the following question:\n\nDocument Text:\n{document_text}\n\nQuestion: {query}"

# Function to generate text-to-image from a description
def generate_image_from_text(description):
    # Here you would call a real text-to-image generation API like DALL-E or Stable Diffusion.
    # For now, let's create a simple image with the description text on it as a placeholder.

    # Create a blank image with white background
    img = Image.new('RGB', (500, 300), color=(255, 255, 255))
    
    # Initialize the drawing context
    draw = ImageDraw.Draw(img)
    
    # Set the font (You can choose a font of your choice)
    font = ImageFont.load_default()  # You can replace with a custom font path

    # Add text to the image (description)
    text = description
    draw.text((10, 150), text, fill=(0, 0, 0), font=font)

    # Save image in memory
    img_io = io.BytesIO()
    img.save(img_io, 'PNG')
    img_io.seek(0)  # Go to the beginning of the file

    return img_io  # Return the image as a BytesIO object

# Streamlit interface
st.title("📄 Document Question Answering and Code Documentation")
st.write(
    "Choose one of the following options:\n\n"
    "- Upload a document to ask a question about it.\n"
    "- Provide a code snippet to generate documentation for it."
)

# Ask user what they want to do
user_choice = st.radio(
    "Select the functionality you want to use:",
    ("Upload Document for Q&A", "Provide Code for Documentation", "Text to Image")
)

# Ask user for their Gemini API key
gemini_api_key = st.text_input("Gemini API Key", type="password")
if not gemini_api_key:
    st.info("Please add your Gemini API key to continue.", icon="🗝️")
else:
    # Create a Gemini client
    client = genai.Client(api_key=gemini_api_key)

    # Conditional section based on the user's choice
    if user_choice == "Upload Document for Q&A":
        # Let the user upload a document
        uploaded_file = st.file_uploader("Upload a document (.txt, .md, .pdf, .docx)", type=("txt", "md", "pdf", "docx"))
        question = st.text_area("Now ask a question about the document!", placeholder="Can you give me a short summary?", disabled=not uploaded_file)

        if uploaded_file and question:
            try:
                # Process the uploaded file
                document = uploaded_file.read()
                document_text = ""
                if uploaded_file.type == "application/pdf":
                    from PyPDF2 import PdfReader
                    reader = PdfReader(BytesIO(document))
                    for page in reader.pages:
                        document_text += page.extract_text()
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    from docx import Document
                    doc = Document(BytesIO(document))
                    for para in doc.paragraphs:
                        document_text += para.text
                elif uploaded_file.type == "text/plain" or uploaded_file.type == "text/markdown":
                    document_text = document.decode("utf-8")

                # Prepare the RAG prompt
                rag_prompt_content = rag_prompt(question, document_text)
                persona_description = "You are a helpful assistant."
                user_query = set_persona(persona_description, rag_prompt_content)

                # Call Gemini API with RAG-based query
                response = client.models.generate_content(
                    model="gemini-2.0-flash",
                    contents=[{"parts": [{"text": user_query}]}]
                )

                if response.candidates:
                    answer = response.candidates[0].content.parts[0].text
                    st.write(answer)
                else:
                    st.error("No response from the model.")
            except Exception as e:
                st.error(f"An error occurred while processing the document: {str(e)}")

    elif user_choice == "Provide Code for Documentation":
        # Handle the code input for generating documentation
        code_input = st.text_area("Paste your code here to generate documentation.", placeholder="def add(a, b):\n    return a + b", height=200)
        generate_code_doc = st.button("Generate Documentation for Code")

        if code_input and generate_code_doc:
            try:
                # Prepare persona and few-shot examples for code documentation
                persona_description = "You are a technical writer who provides concise and clear explanations."
                user_query = set_persona(persona_description, code_input)
                few_shot_examples = [
                    "Example 1: Code: def add(a, b): return a + b. Documentation: This function adds two numbers and returns the result.",
                    "Example 2: Code: def subtract(a, b): return a - b. Documentation: This function subtracts the second argument from the first and returns the result."
                ]
                prompt = get_few_shot_prompt(user_query, few_shot_examples)

                # Generate documentation using the Gemini API
                response_code = client.models.generate_content(
                    model="gemini-2.0-flash",
                    contents=[{"parts": [{"text": prompt}]}]
                )

                if response_code.candidates:
                    doc_answer = response_code.candidates[0].content.parts[0].text
                    st.write(doc_answer)

                    # Create a Word document with the generated documentation
                    doc = Document()
                    doc.add_heading('Code Documentation', 0)
                    add_styled_text(doc, "Overview", style="Heading 1", is_bold=True)
                    add_styled_text(doc, "This document provides a detailed explanation of the code snippet provided by the user.", is_italic=True)
                    doc.add_paragraph("\n")

                    parse_markdown_to_word(doc, doc_answer)
                    add_tip_or_warning_box(doc, "Remember to validate the code syntax before running it.", box_type="tip")

                    # Save the document in memory
                    doc_io = BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)

                    # Provide download link for the Word document
                    st.download_button(
                        label="Download Documentation as Word File",
                        data=doc_io,
                        file_name="code_documentation_confluence_style.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("No response from the model.")
            except Exception as e:
                st.error(f"An error occurred while generating code documentation: {str(e)}")

    elif user_choice == "Text to Image":
        description = st.text_area("Enter description for the image:")
        if description:
          image = generate_image_from_text(description)  # Generate image from description
          st.image(image, caption="Generated Image")  # Display the generated image
