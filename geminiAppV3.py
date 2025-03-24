import streamlit as st
from google import genai
from io import BytesIO
from docx import Document
import re  # For markdown-like syntax processing

# Function to add styled text in the Word document
def add_styled_text(doc, text, style=None, is_bold=False):
    """
    Adds a paragraph with optional style to the Word document. Handles bold text within a paragraph.
    """
    para = doc.add_paragraph()
    if style:
        para.style = style
    if is_bold:
        run = para.add_run(text)
        run.bold = True
    else:
        para.add_run(text)

# Function to convert markdown-like syntax to Word formatting
def parse_markdown_to_word(doc, text):
    """
    Converts markdown-like syntax to Word document styles (e.g., **bold**, # Heading).
    """
    # Convert headers (e.g., # Header) to Heading 1 style
    text = re.sub(r'^\#\s+(.*)', lambda m: add_styled_text(doc, m.group(1), style='Heading 1'), text, flags=re.MULTILINE)
    
    # Convert bold text (e.g., **bold**) to actual bold
    # Find all bold text using regex and handle it
    while '**' in text:
        start = text.find('**')
        end = text.find('**', start + 2)
        if start == -1 or end == -1:
            break
        bold_text = text[start + 2:end]
        # Add the portion before the bold text
        add_styled_text(doc, text[:start])
        # Add the bold text
        add_styled_text(doc, bold_text, is_bold=True)
        # Move the text past the bold
        text = text[end + 2:]
    
    # After all bold text is processed, add the remaining normal text
    if text:
        add_styled_text(doc, text)

# Show title and description.
st.title("📄 Document Question Answering and Code Documentation")
st.write(
    "Upload a document below and ask a question about it, or provide a code snippet to generate documentation."
    " To use this app, you need to provide your Gemini API key, which you can get from the Gemini platform."
)

# Ask user for their Gemini API key via `st.text_input`.
gemini_api_key = st.text_input("Gemini API Key", type="password")
if not gemini_api_key:
    st.info("Please add your Gemini API key to continue.", icon="🗝️")
else:
    # Create a Gemini client.
    client = genai.Client(api_key=gemini_api_key)

    # Let the user upload a document or provide code.
    uploaded_file = st.file_uploader(
        "Upload a document (.txt, .md, .pdf, .docx)", type=("txt", "md", "pdf", "docx")
    )
    
    code_input = st.text_area(
        "Or provide a code snippet here to generate documentation for it.",
        placeholder="def add(a, b):\n    return a + b",
        height=200
    )
    
    # Ask the user for a question via `st.text_area`.
    question = st.text_area(
        "Now ask a question about the document!",
        placeholder="Can you give me a short summary?",
        disabled=not uploaded_file and not code_input,
    )

    # Option to generate documentation for the provided code
    generate_code_doc = st.button("Generate Documentation for Code")

    # If there is an uploaded document or code snippet, process it
    if uploaded_file and question:
        try:
            # Process the uploaded file and question.
            document = uploaded_file.read()

            # Handle different file types (PDF, DOCX)
            if uploaded_file.type == "application/pdf":
                from PyPDF2 import PdfReader
                reader = PdfReader(BytesIO(document))
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                document = text

            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                from docx import Document
                doc = Document(BytesIO(document))
                text = ""
                for para in doc.paragraphs:
                    text += para.text
                document = text

            # Prepare the contents to send to Gemini API for the document question-answering
            content = f"Here's a document: {document} \n\n---\n\n {question}"

            # Generate an answer using the Gemini API for document
            response = client.models.generate_content(
                model="gemini-2.0-flash", 
                contents=[{"parts": [{"text": content}]}]
            )

            # Access the first candidate and its text
            if response.candidates:
                answer = response.candidates[0].content.parts[0].text
                st.write(answer)
            else:
                st.error("No response from the model.")
                
        except Exception as e:
            st.error(f"An error occurred while processing the document: {str(e)}")

    # Handle the code input for generating documentation
    if code_input and generate_code_doc:
        try:
            # Prepare the contents to send to Gemini API for generating code documentation
            code_content = f"Here's a code snippet: {code_input} \n\n---\n\n Can you generate documentation for this code?"

            # Generate documentation using the Gemini API
            response_code = client.models.generate_content(
                model="gemini-2.0-flash", 
                contents=[{"parts": [{"text": code_content}]}]
            )

            if response_code.candidates:
                doc_answer = response_code.candidates[0].content.parts[0].text
                st.write(doc_answer)
                
                # Create a Word document with the generated documentation
                doc = Document()
                doc.add_heading('Code Documentation', 0)

                # Parse the generated documentation and add formatted text
                parse_markdown_to_word(doc, doc_answer)

                # Save the document in memory
                doc_io = BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)

                # Provide download link for the Word document
                st.download_button(
                    label="Download Documentation as Word File",
                    data=doc_io,
                    file_name="code_documentation.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            else:
                st.error("No response from the model.")
                
        except Exception as e:
            st.error(f"An error occurred while generating code documentation: {str(e)}")
