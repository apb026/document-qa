import streamlit as st
from google import genai
from io import BytesIO
from docx import Document
import re

# Function to add styled text in the Word document
def add_styled_text(doc, text, style=None, is_bold=False, is_italic=False, is_code=False):
    """
    Adds a paragraph with optional style to the Word document. Handles bold text within a paragraph.
    """
    para = doc.add_paragraph()
    
    if is_code:
        run = para.add_run(text)
        run.font.name = 'Courier New'  # Use monospace font for code
        para.style = 'Normal'
    else:
        if style:
            para.style = style
        run = para.add_run(text)
        
        # Apply bold or italic formatting if needed
        if is_bold:
            run.bold = True
        if is_italic:
            run.italic = True

# Function to convert markdown-like syntax to Word formatting
def parse_markdown_to_word(doc, text):
    """
    Converts markdown-like syntax to Word document styles (e.g., **bold**, # Heading).
    """
    # Convert headers (e.g., # Header) to Heading 1 style
    text = re.sub(r'^\#\s+(.*)', lambda m: add_styled_text(doc, m.group(1), style='Heading 1'), text, flags=re.MULTILINE)
    
    # Convert bold text (e.g., **bold**) to actual bold
    while '**' in text:
        start = text.find('**')
        end = text.find('**', start + 2)
        if start == -1 or end == -1:
            break
        bold_text = text[start + 2:end]
        # Add the portion before the bold text
        add_styled_text(doc, text[:start])
        # Add the bold text
        add_styled_text(doc, bold_text, is_bold=True)
        # Move the text past the bold
        text = text[end + 2:]
    
    # Add any remaining text
    if text:
        add_styled_text(doc, text)

# Function to add a tip or warning box (like Confluence)
def add_tip_or_warning_box(doc, text, box_type='tip'):
    """
    Add a tip or warning box to the Word document
    """
    # Create a shaded box with tip or warning
    para = doc.add_paragraph()
    para.add_run(text).italic = True
    
    if box_type == 'tip':
        # For tips, we can use a light blue background (mimicking Confluence)
        para.paragraph_format.alignment = 1  # Center-aligned
        doc.add_paragraph("TIP:", style='Heading 3')
        para = doc.add_paragraph()
        para.add_run(text).bold = True
    elif box_type == 'warning':
        para.paragraph_format.alignment = 1
        doc.add_paragraph("WARNING:", style='Heading 3')
        para = doc.add_paragraph()
        para.add_run(text).bold = True

# Show title and description.
st.title("📄 Document Question Answering and Code Documentation")
st.write(
    "Choose one of the following options:\n\n"
    "- Upload a document to ask a question about it.\n"
    "- Provide a code snippet to generate documentation for it."
)

# Ask user what they want to do
user_choice = st.radio(
    "Select the functionality you want to use:",
    ("Upload Document for Q&A", "Provide Code for Documentation")
)

# Ask user for their Gemini API key via `st.text_input`.
gemini_api_key = st.text_input("Gemini API Key", type="password")
if not gemini_api_key:
    st.info("Please add your Gemini API key to continue.", icon="🗝️")
else:
    # Create a Gemini client.
    client = genai.Client(api_key=gemini_api_key)

    # Conditional section based on the user's choice
    if user_choice == "Upload Document for Q&A":
        # Let the user upload a document.
        uploaded_file = st.file_uploader(
            "Upload a document (.txt, .md, .pdf, .docx)", type=("txt", "md", "pdf", "docx")
        )
        
        # Ask the user for a question via `st.text_area`.
        question = st.text_area(
            "Now ask a question about the document!",
            placeholder="Can you give me a short summary?",
            disabled=not uploaded_file,
        )

        # If the document is uploaded and question is provided, process it
        if uploaded_file and question:
            try:
                # Process the uploaded file and question.
                document = uploaded_file.read()

                # Handle different file types (PDF, DOCX)
                if uploaded_file.type == "application/pdf":
                    from PyPDF2 import PdfReader
                    reader = PdfReader(BytesIO(document))
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text()
                    document = text

                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    from docx import Document
                    doc = Document(BytesIO(document))
                    text = ""
                    for para in doc.paragraphs:
                        text += para.text
                    document = text

                # Prepare the contents to send to Gemini API for the document question-answering
                content = f"Here's a document: {document} \n\n---\n\n {question}"

                # Generate an answer using the Gemini API for document
                response = client.models.generate_content(
                    model="gemini-2.0-flash", 
                    contents=[{"parts": [{"text": content}]}]
                )

                # Access the first candidate and its text
                if response.candidates:
                    answer = response.candidates[0].content.parts[0].text
                    st.write(answer)
                else:
                    st.error("No response from the model.")
                
            except Exception as e:
                st.error(f"An error occurred while processing the document: {str(e)}")

    elif user_choice == "Provide Code for Documentation":
        # Handle the code input for generating documentation
        code_input = st.text_area(
            "Paste your code here to generate documentation.",
            placeholder="def add(a, b):\n    return a + b",
            height=200
        )
        
        generate_code_doc = st.button("Generate Documentation for Code")

        # If there is code and button is pressed, generate documentation
        if code_input and generate_code_doc:
            try:
                # Prepare the contents to send to Gemini API for generating code documentation
                code_content = f"Here's a code snippet: {code_input} \n\n---\n\n Can you generate documentation for this code?"

                # Generate documentation using the Gemini API
                response_code = client.models.generate_content(
                    model="gemini-2.0-flash", 
                    contents=[{"parts": [{"text": code_content}]}]
                )

                if response_code.candidates:
                    doc_answer = response_code.candidates[0].content.parts[0].text
                    st.write(doc_answer)
                    
                    # Create a Word document with the generated documentation
                    doc = Document()
                    doc.add_heading('Code Documentation', 0)

                    # Add introductory sections
                    add_styled_text(doc, "Overview", style="Heading 1", is_bold=True)
                    add_styled_text(doc, "This document provides a detailed explanation of the code snippet provided by the user.", is_italic=True)
                    doc.add_paragraph("\n")

                    # Parse the generated documentation and add formatted text
                    parse_markdown_to_word(doc, doc_answer)

                    # Optionally, add a Tip Box or Warning Box
                    add_tip_or_warning_box(doc, "Remember to validate the code syntax before running it.", box_type="tip")

                    # Save the document in memory
                    doc_io = BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)

                    # Provide download link for the Word document
                    st.download_button(
                        label="Download Documentation as Word File",
                        data=doc_io,
                        file_name="code_documentation_confluence_style.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                else:
                    st.error("No response from the model.")
                
            except Exception as e:
                st.error(f"An error occurred while generating code documentation: {str(e)}")